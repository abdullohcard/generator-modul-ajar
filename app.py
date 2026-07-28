import streamlit as st
import google.generativeai as genai
import time
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

st.set_page_config(page_title="Generator Modul Ajar Deep Learning", page_icon="📚", layout="wide")

st.title("📚 Generator Modul Ajar Deep Learning (Semua Mapel)")
st.subheader("SMP Tri Sukses Boarding School Kota Jambi (1 JP = 30 Menit)")
st.caption("Berbasis CP Terbaru BSKAP 046/H/KR/2025 & Kepka BKPDM No. 020 Tahun 2026 - Kurikulum Merdeka")

# FUNGSI PEMPROSES FORMAT TEKS KE WORD (MARKDOWN TO DOCX)
def add_formatted_runs(paragraph, text, base_bold=False, font_size=10.5):
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for token in tokens:
        if not token:
            continue
        
        is_bold = base_bold
        is_italic = False
        clean_text = token

        if token.startswith('**') and token.endswith('**'):
            is_bold = True
            clean_text = token[2:-2]
        elif token.startswith('*') and token.endswith('*'):
            is_italic = True
            clean_text = token[1:-1]

        run = paragraph.add_run(clean_text)
        run.font.name = 'Arial'
        run.font.size = Pt(font_size)
        run.font.bold = is_bold
        run.font.italic = is_italic

def generate_docx(title, content):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"MODUL AJAR DEEP LEARNING\n{title.upper()}")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(14)
    run_title.font.bold = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    lines = content.split('\n')
    for line in lines:
        line_strip = line.strip()
        if not line_strip or line_strip == '---':
            continue
            
        if line_strip.startswith('# '):
            p = doc.add_paragraph()
            add_formatted_runs(p, line_strip[2:], base_bold=True, font_size=12)
        elif line_strip.startswith('## '):
            p = doc.add_paragraph()
            add_formatted_runs(p, line_strip[3:], base_bold=True, font_size=11)
        elif line_strip.startswith('### '):
            p = doc.add_paragraph()
            add_formatted_runs(p, line_strip[4:], base_bold=True, font_size=10.5)
        elif line_strip.startswith('* ') or line_strip.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_runs(p, line_strip[2:], font_size=10.5)
        else:
            p = doc.add_paragraph()
            add_formatted_runs(p, line_strip, font_size=10.5)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# BACA API KEY DARI SECRETS
api_keys = []
if "GEMINI_API_KEYS" in st.secrets:
    api_keys = [k.strip() for k in st.secrets["GEMINI_API_KEYS"].split(",") if k.strip()]
elif "GEMINI_API_KEY" in st.secrets:
    api_keys = [st.secrets["GEMINI_API_KEY"]]

# Sidebar
st.sidebar.header("⚙️ Pengaturan Aplikasi")
if api_keys:
    st.sidebar.success(f"✅ Akses AI Otomatis Aktif!")
else:
    manual_key = st.sidebar.text_input("Masukkan Gemini API Key (Manual):", type="password")
    if manual_key:
        api_keys = [manual_key]

st.sidebar.markdown("---")
st.sidebar.info("""
**Aturan Standar Sekolah:**
- 1 JP = 30 Menit
- Strategi Deep Learning: Mindful, Meaningful, Joyful
- Pendekatan Ramah Asrama / Non-Gadget (Unplugged)
- Acuan: BSKAP 046/H/KR/2025 & Kepka BKPDM No. 020 Tahun 2026
""")

# DATABASE MULTI-MAPEL DENGAN ELEMEN & REKOMENDASI TOPIK CP TERBARU
DATABASE_MAPEL = {
    "Matematika": {
        "Bilangan": [
            "Operasi Aritmatika Bilangan Bulat dan Rasional dalam Kehidupan Asrama",
            "Penerapan Rasio, Skala, dan Proporsi untuk Pembagian Tugas Piket",
            "Lainnya (Ketik Manual)"
        ],
        "Aljabar": [
            "Penyelesaian Persamaan dan Pertidaksamaan Linear Satu Variabel",
            "Memahami Relasi, Fungsi, dan Grafik dalam Kehidupan Sehari-hari",
            "Lainnya (Ketik Manual)"
        ],
        "Pengukuran & Geometri": [
            "Menghitung Luas Permukaan dan Volume Bangun Ruang di Lingkungan Sekolah",
            "Teorema Pythagoras dan Transformasi Geometri Sederhana",
            "Lainnya (Ketik Manual)"
        ],
        "Analisis Data dan Peluang": [
            "Penyajian dan Analisis Data Pengamatan Santri Menggunakan Diagram",
            "Konsep Peluang Sederhana dan Frekuensi Harapan",
            "Lainnya (Ketik Manual)"
        ]
    },
    "Bahasa Indonesia": {
        "Menyimak": [
            "Menganalisis Gagasan dan Pesan dari Teks Lisan Nonsastra",
            "Menganalisis Unsur Intrinsik Cerita Pendek/Kisah Inspiratif",
            "Lainnya (Ketik Manual)"
        ],
        "Membaca dan Memirsa": [
            "Menganalisis Informasi dan Menemukan Makna Tersirat Teks Multimodal",
            "Evaluasi Kualitas dan Kredibilitas Bacaan Digital/Cetak",
            "Lainnya (Ketik Manual)"
        ],
        "Berbicara dan Mempresentasikan": [
            "Mempresentasikan Gagasan dan Solusi Masalah Secara Kritis dan Santun",
            "Diskusi dan Dialog Logis Terkait Isu Lingkungan Asrama",
            "Lainnya (Ketik Manual)"
        ],
        "Menulis": [
            "Menulis Teks Laporan Hasil Observasi (LHO) Lingkungan Sekolah",
            "Menulis Teks Prosedur/Eksposisi Berbasis Pengalaman Santri",
            "Lainnya (Ketik Manual)"
        ]
    },
    "Bahasa Inggris": {
        "Listening - Speaking": [
            "Expressing Ideas and Opinions about Daily Routine in Boarding School",
            "Asking and Giving Directions/Information in School Context",
            "Lainnya (Ketik Manual)"
        ],
        "Reading - Viewing": [
            "Understanding Main Ideas and Detailed Information from Descriptive Text",
            "Analyzing Explicit and Implicit Information in Narrative Text",
            "Lainnya (Ketik Manual)"
        ],
        "Writing - Presenting": [
            "Writing Simple Procedure Text / Recount Text about Personal Experience",
            "Presenting Group Project Results using Simple Compound Sentences",
            "Lainnya (Ketik Manual)"
        ]
    },
    "Pendidikan Pancasila": {
        "Pancasila": [
            "Sejarah Kelahiran Pancasila dan Penerapan Nilainya di Asrama",
            "Kedudukan Pancasila sebagai Dasar Negara dan Ideologi Bangsa",
            "Lainnya (Ketik Manual)"
        ],
        "UUD 1945": [
            "Penerapan Hak dan Kewajiban Santri sebagai Warga Sekolah dan Negara",
            "Musyawarah Mufakat dalam Membuat Kesepakatan Bersama di Kelas/Asrama",
            "Lainnya (Ketik Manual)"
        ],
        "Bhinneka Tunggal Ika": [
            "Mengembangkan Toleransi dan Menghargai Keragaman Suku/Budaya Santri",
            "Pelestarian Tradisi dan Kearifan Lokal Jambi",
            "Lainnya (Ketik Manual)"
        ],
        "NKRI": [
            "Peran Santri dalam Menjaga Keutuhan Wilayah NKRI dan Bela Negara",
            "Lainnya (Ketik Manual)"
        ]
    },
    "IPA": {
        "Pemahaman IPA": [
            "Klasifikasi Makhluk Hidup dan Sistem Organisasi Kehidupan",
            "Interaksi Antar Makhluk Hidup, Ekosistem, dan Mitigasi Perubahan Iklim",
            "Ragam Gerak, Gaya, Tekanan, serta Hubungan Usaha dan Energi",
            "Suhu, Kalor, Gelombang, Kemagnetan, dan Kelistrikan Ramah Lingkungan",
            "Posisi Bumi-Bulan-Matahari dan Tata Surya",
            "Lainnya (Ketik Manual)"
        ],
        "Keterampilan Proses": [
            "Penyelidikan Ilmiah Sederhana: Mengamati, Menganalisis Data, dan Komunikasi",
            "Lainnya (Ketik Manual)"
        ]
    },
    "IPS": {
        "Pemahaman Konsep": [
            "Kehidupan Sosial dan Kondisi Lingkungan Sekitar Santri",
            "Konektivitas Antarruang dan Kegiatan Ekonomi Masyarakat",
            "Interaksi Sosial, Lembaga Sosial, dan Dinamika Sosial Budaya",
            "Sejarah Lokal, Asal-Usul Nenek Moyang, dan Jalur Rempah Nusantara",
            "Lainnya (Ketik Manual)"
        ],
        "Keterampilan Proses": [
            "Investasi/Observasi Sederhana Masalah Sosial dan Lingkungan",
            "Lainnya (Ketik Manual)"
        ]
    },
    "Informatika": {
        "Berpikir Komputasional": [
            "Penerapan BK untuk Pemecahan Masalah Sehari-hari Santri di Asrama",
            "Pengenalan Konsep Himpunan Data Terstruktur dalam Kehidupan Sekolah",
            "Penulisan Sekumpulan Instruksi Algoritma Menggunakan Pseudocode Sederhana",
            "Lainnya (Ketik Manual)"
        ],
        "Literasi Digital": [
            "Cara Kerja Mesin Pencari dan Evaluasi Kredibilitas Informasi (Anti-Hoaks)",
            "Pemanfaatan Pengolah Dokumen, Lembar Kerja, dan Presentasi",
            "Komponen, Fungsi Utama Komputer, dan Jaringan Lokal/Internet",
            "Rekam Jejak Digital, Etika Komunikasi, dan Keamanan Data Pribadi",
            "Lainnya (Ketik Manual)"
        ],
        "Analisis Data & Algoritma": [
            "Pemrosesan Data Sederhana dan Logika Program Unplugged",
            "Lainnya (Ketik Manual)"
        ]
    },
    "Koding dan Kecerdasan Artifisial": {
        "Konsep Dasar Koding (Programming)": [
            "Pengenalan Logika Pemrograman Block-Based / Unplugged Flowchart",
            "Struktur Kontrol: Urutan (Sequencing), Percabangan (If-Else), dan Perulangan (Loop)",
            "Lainnya (Ketik Manual)"
        ],
        "Kecerdasan Artifisial (AI) & Etika": [
            "Pengenalan Konsep Dasar AI, Machine Learning, dan Computer Vision",
            "Pemanfaatan AI Secara Bijak, Aman, dan Etis dalam Pendidikan Santri",
            "Lainnya (Ketik Manual)"
        ]
    },
    "Seni Rupa": {
        "Mengalami & Menciptakan": [
            "Menganalisis Unsur Rupa dan Prinsip Desain pada Karya Seni",
            "Membuat Karya Seni Rupa 2D/3D Menggunakan Bahan Alam Sekitar Asrama",
            "Lainnya (Ketik Manual)"
        ],
        "Merefleksikan & Berdampak": [
            "Apresiasi Karya Seni Rupa Teman dan Pengaruhnya Terhadap Lingkungan",
            "Lainnya (Ketik Manual)"
        ]
    },
    "PJOK": {
        "Keterampilan & Pengetahuan Gerak": [
            "Praktik dan Analisis Variasi Pola Gerak Dasar Permainan Bola Besar/Kecil",
            "Aktivitas Kebugaran Jasmani untuk Menjaga Kesehatan Santri di Asrama",
            "Lainnya (Ketik Manual)"
        ],
        "Perilaku Hidup Sehat": [
            "Pencegahan Pergaulan Bebas, Zat Aditif/Adiktif, dan Pola Makan Sehat",
            "Lainnya (Ketik Manual)"
        ]
    },
    "Bimbingan Konseling (BK)": {
        "Bimbingan Pribadi & Sosial": [
            "Pengembangan Adaptasi Diri, Kemandirian, dan Manajemen Waktu di Asrama",
            "Pengelolaan Emosi, Hubungan Pertemanan Sehat, dan Anti-Bullying",
            "Lainnya (Ketik Manual)"
        ],
        "Bimbingan Belajar & Karir": [
            "Strategi Belajar Efektif (Growth Mindset) dan Pengenalan Minat/Bakat",
            "Lainnya (Ketik Manual)"
        ]
    },
    "PAI-BP": {
        "Al-Qur'an Hadis": [
            "Membaca, Menghafal, Menulis, dan Menjelaskan Ayat/Hadis tentang Iman, Takwa, dan Toleransi",
            "Lainnya (Ketik Manual)"
        ],
        "Akidah & Akhlak": [
            "Meyakini Rukun Iman serta Menerapkan Sikap Ikhlas, Bersyukur, dan Khusnuzan",
            "Menerapkan Cinta Allah/Rasul, Cinta Ilmu, Cinta Lingkungan, dan Diri Sendiri",
            "Lainnya (Ketik Manual)"
        ],
        "Fikih & Sejarah Peradaban Islam": [
            "Ketentuan Sujud, Rukhsah Salat, Ketentuan Salat Jumat, Puasa, dan Bersuci",
            "Kisah Keteladanan Nabi Muhammad SAW, Khulafaur Rasyidin, dan Masuknya Islam di Indonesia",
            "Lainnya (Ketik Manual)"
        ]
    }
}

# Form Input Data Modul
st.header("📝 Form Isian Modul Ajar")

col1, col2 = st.columns(2)

with col1:
    nama_guru = st.text_input("Nama Guru Penyusun:", value="Muhammad Irfa'udin Aulia, S.Kom., Gr.")
    
    # DROPDOWN PILIHAN MATA PELAJARAN (LENGKAP)
    mapel = st.selectbox("Mata Pelajaran", list(DATABASE_MAPEL.keys()))
    kelas = st.selectbox("Kelas (Fase D)", ["Kelas 7", "Kelas 8", "Kelas 9"])
    semester = st.selectbox("Semester", ["Ganjil", "Genap"])
    
    # DROPDOWN ELEMEN CP DINAMIS SESUAI MAPEL
    elemen_cp = st.selectbox("Elemen CP", list(DATABASE_MAPEL[mapel].keys()))
    
    # DROPDOWN REKOMENDASI TOPIK DINAMIS SESUAI ELEMEN MAPEL
    pilihan_topik = st.selectbox("💡 Pilih Rekomendasi Topik/Materi (Sesuai CP Terbaru):", DATABASE_MAPEL[mapel][elemen_cp])
    
    if pilihan_topik == "Lainnya (Ketik Manual)":
        topik_final = st.text_input("Ketikkan Topik/Pokok Bahasan Custom:", placeholder="Contoh: Pembiasaan Karakter Santri")
    else:
        topik_final = pilihan_topik

with col2:
    tahun_ajaran = st.text_input("Tahun Pelajaran:", value="2026 / 2027")
    jumlah_pertemuan = st.number_input("Jumlah Pertemuan", min_value=1, max_value=5, value=2)
    jp_per_pertemuan = st.number_input("Alokasi JP per Pertemuan (1 JP = 30 mnt)", min_value=1, max_value=4, value=2)
    pendekatan = st.radio("Pendekatan Pembelajaran", ["Unplugged (Tanpa Gadget/Cetak/Papan Tulis)", "Plugged (Praktik Lab/Komputer/Lapangan)"])

st.markdown("---")

if st.button("🚀 Buat Modul Ajar Sesuai Templat", type="primary"):
    if not api_keys:
        st.error("API Key belum dikonfigurasi.")
    elif not topik_final:
        st.warning("Pilih atau ketikkan Topik terlebih dahulu.")
    else:
        total_jp = jumlah_pertemuan * jp_per_pertemuan
        total_menit = jp_per_pertemuan * 30
        
        prompt = f"""
        Anda adalah konsultan kurikulum Kurikulum Merdeka Kemendikdasmen Indonesia.
        Buatkan Modul Ajar Pembelajaran Mendalam (Deep Learning) secara SANGAT LENGKAP DAN DETAIL dengan FORMAT WAKTU DAN STRUKTUR PERSIS SESUAI TEMPLAT DOKUMEN BERIKUT.

        === DATA INPUT MODUL ===
        - Nama Sekolah: SMP Tri Sukses Boarding School Kota Jambi
        - Nama Penyusun: {nama_guru}
        - Mata Pelajaran: {mapel}
        - Fase / Kelas / Semester: D / {kelas} / {semester}
        - Alokasi Waktu: {total_jp} Jam Pelajaran ({jumlah_pertemuan} pertemuan x {jp_per_pertemuan} JP)
        - Tahun Pelajaran: {tahun_ajaran}
        - Elemen CP: {elemen_cp} (Berdasarkan Keputusan BSKAP No. 046/H/KR/2025 & Kepka BKPDM No. 020 Tahun 2026)
        - Topik / Bab: {topik_final}
        - Pendekatan: {pendekatan}
        - Aturan Durasi: 1 JP = 30 Menit (Setiap pertemuan = {total_menit} Menit).

        === STRUKTUR OUTPUT WAJIB (JANGAN MENGUBAH JUDUL SEKSI) ===

        MODUL AJAR DEEP LEARNING
        MAPEL : {mapel}
        Bab / Topik : {topik_final}

        A. Identitas Modul
        - Nama Sekolah : SMP Tri Sukses Boarding School Kota Jambi
        - Nama Penyusun : {nama_guru}
        - Mata Pelajaran : {mapel}
        - Fase / Kelas / Semester : D / {kelas} / {semester}
        - Alokasi Waktu : {total_jp} Jam Pelajaran ({jumlah_pertemuan} pertemuan x {jp_per_pertemuan} JP)
        - Tahun Pelajaran : {tahun_ajaran}

        B. Identifikasi Kesiapan Peserta Didik
        (Tuliskan 2-3 kalimat penjelasan kesiapan awal santri/peserta didik kelas {kelas} terkait materi {topik_final}).

        C. Karakteristik Materi Pelajaran
        (Tuliskan deskripsi ringkas karakteristik materi {topik_final} dalam pembelajaran mendalam).

        D. Dimensi Profil Lulusan Pembelajaran
        Dalam pembelajaran ini, dimensi profil lulusan yang akan dicapai adalah:
        - Kewargaan: ...
        - Penalaran Kritis: ...
        - Kolaborasi: ...
        - Kemandirian: ...
        - Komunikasi: ...

        DESAIN PEMBELAJARAN
        A. Capaian Pembelajaran (CP) Nomor : BSKAP No. 046/H/KR/2025 jo. Kepka BKPDM No. 020 Tahun 2026
        Pada akhir fase D, peserta didik diharapkan mampu:
        - Pengetahuan: ...
        - Keterampilan: ...
        - Sikap: ...

        B. Lintas Disiplin Ilmu yang Relevan
        - Bahasa Indonesia: ...
        - Seni Budaya: ...
        - Matematika: ...
        - Pendidikan Pancasila: ...
        - Ilmu Pengetahuan Alam (IPA): ...

        C. Tujuan Pembelajaran
        (Rincikan Tujuan Pembelajaran khusus untuk Pertemuan 1 sampai Pertemuan {jumlah_pertemuan} secara operasional dan terukur).

        D. Topik Pembelajaran Kontekstual
        Topik pembelajaran akan berpusat pada lingkungan sekitar peserta didik/kehidupan asrama. Berikan 3-4 contoh pertanyaan kontekstual.

        E. Kerangka Pembelajaran
        - Praktik Pedagogik: (Metode Berbasis Proyek/Diskusi/Simulasi/Wawancara/Praktik Lapangan)
        - Mitra Pembelajaran: (Lingkungan Sekolah, Luar Sekolah, Masyarakat/Asrama)
        - Lingkungan Belajar: (Ruang Fisik/Lapangan/Lab, Ruang Virtual, Budaya Belajar)
        - Pemanfaatan Digital: (Perencanaan & Pelaksanaan)

        F. Langkah-Langkah Pembelajaran Berdiferensiasi
        (Buat rinci untuk PERTEMUAN 1 hingga PERTEMUAN {jumlah_pertemuan}. Setiap pertemuan WAJIB membagi alokasi {total_menit} Menit dengan alur 3M berikut:)

        Pertemuan [X]: [Judul Sub-Topik Pertemuan]
        1. Kegiatan Pendahuluan ({int(total_menit*0.15)} Menit)
           - Berkesadaran (Mindful): (Sertakan instruksi apersepsi/ice breaking & pertanyaan pemantik)
           - Bermakna (Meaningful): (Kaitkan dengan kehidupan nyata santri di asrama)
           - Menggembirakan (Joyful): (Sertakan permainan/simulasi singkat)
        2. Kegiatan Inti ({int(total_menit*0.70)} Menit)
           - Memahami (Diferensiasi Konten): (Eksplorasi konsep dalam berbagai format/metode)
           - Mengaplikasi (Diferensiasi Proses): (Kerja kelompok/LKPD/studi kasus kontekstual/praktik)
           - Merefleksi (Berkesadaran, Bermakna): (Presentasi, pemaknaan & komitmen positif)
        3. Kegiatan Penutup ({int(total_menit*0.15)} Menit)
           - (Umpan balik, kesimpulan bersama, pengantar pertemuan berikutnya)

        G. Asesmen Pembelajaran
        1. Asesmen Awal Pembelajaran:
           - Observasi: ...
           - Kuesioner / Tes Diagnostik: ...
           - Soal Asesmen Awal: (Tuliskan 5 Soal Asesmen Awal LENGKAP dengan teks soalnya)
        2. Asesmen Proses Pembelajaran:
           - Tugas Harian & Diskusi Kelompok: ...
           - Soal / Instrumen Asesmen Proses: (Tuliskan 5 Pertanyaan/Tugas Proses LENGKAP)
        3. Asesmen Akhir Pembelajaran:
           - Jurnal Reflektif & Tugas Akhir/Proyek: ...
           - Soal Asesmen Akhir: (Tuliskan 5 Soal Uraian Akhir LENGKAP)
        """

        success = False
        response_text = ""
        last_err = ""

        with st.spinner(f"AI sedang menyusun Modul Ajar {mapel} Sesuai Format Templat Resmi... Mohon tunggu sebentar."):
            for key in api_keys:
                if success:
                    break
                try:
                    genai.configure(api_key=key)
                    active_models = [
                        m.name for m in genai.list_models() 
                        if 'generateContent' in m.supported_generation_methods
                    ]
                    
                    for m_name in active_models:
                        try:
                            model = genai.GenerativeModel(m_name)
                            res = model.generate_content(prompt)
                            response_text = res.text
                            success = True
                            break
                        except Exception as e_mod:
                            last_err = str(e_mod)
                            continue
                except Exception as e_key:
                    last_err = str(e_key)
                    continue

        if success:
            st.success(f"✨ Modul Ajar {mapel} Berhasil Dibuat!")
            st.markdown("---")
            st.markdown(response_text)
            st.markdown("---")
            
            col_dl1, col_dl2 = st.columns(2)
            docx_file = generate_docx(f"{mapel} - {topik_final}", response_text)
            
            with col_dl1:
                st.download_button(
                    label="📄 Unduh File MS Word (.docx)",
                    data=docx_file,
                    file_name=f"Modul_Ajar_{mapel}_{kelas}_{topik_final}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
            with col_dl2:
                st.download_button(
                    label="💾 Unduh Teks (.txt)",
                    data=response_text,
                    file_name=f"Modul_Ajar_{mapel}_{kelas}_{topik_final}.txt",
                    mime="text/plain"
                )
        else:
            st.error("Gagal menghubungkan ke AI. Silakan coba lagi dalam beberapa detik.")
            st.caption(f"Error detail: {last_err}")
